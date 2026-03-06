from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.shared import Inches


def export_blog_docx(title: str, sections: list[dict], logo_bytes: bytes | None = None) -> bytes:
    document = Document()

    if logo_bytes:
        try:
            document.add_picture(BytesIO(logo_bytes), width=Inches(1.8))
        except Exception:
            pass

    document.add_heading(title, level=1)

    for section in sections:
        heading = (section.get("heading") or "").strip()
        content = (section.get("content") or "").strip()

        if heading:
            document.add_heading(heading, level=2)

        if content:
            for para in [p.strip() for p in content.split("\n") if p.strip()]:
                document.add_paragraph(para)

    output = BytesIO()
    document.save(output)
    return output.getvalue()
